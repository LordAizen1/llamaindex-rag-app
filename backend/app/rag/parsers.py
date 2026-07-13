"""Per-format document parsers.

Each parser normalizes a file into a list of LlamaIndex `Document`s carrying
metadata that later propagates onto every chunk:
  - PDF  -> one Document per page,   metadata {source, page}
  - DOCX -> one Document per section, metadata {source, section}
  - MD   -> one Document per heading, metadata {source, section}
"""
import os

from docx import Document as DocxDocument
from llama_index.core import Document
from pypdf import PdfReader


class UnsupportedFileType(Exception):
    pass


class EmptyDocument(Exception):
    pass


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}


def _file_type(filename: str) -> str:
    return os.path.splitext(filename)[1].lower().lstrip(".")


def parse_pdf(path: str, filename: str) -> list[Document]:
    reader = PdfReader(path)
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        docs.append(
            Document(
                text=text,
                metadata={"source": filename, "page": i + 1, "file_type": "pdf"},
            )
        )
    return docs


def parse_docx(path: str, filename: str) -> list[Document]:
    doc = DocxDocument(path)
    docs: list[Document] = []
    current_heading: str | None = None
    buf: list[str] = []

    def flush() -> None:
        nonlocal buf
        if buf:
            docs.append(
                Document(
                    text="\n".join(buf).strip(),
                    metadata={
                        "source": filename,
                        "section": current_heading or "Body",
                        "file_type": "docx",
                    },
                )
            )
            buf = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if style.startswith("Heading"):
            flush()
            current_heading = text
        else:
            buf.append(text)
    flush()
    return docs


def parse_markdown(path: str, filename: str) -> list[Document]:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        lines = f.readlines()

    docs: list[Document] = []
    current_heading: str | None = None
    buf: list[str] = []

    def flush() -> None:
        nonlocal buf
        if buf:
            docs.append(
                Document(
                    text="\n".join(buf).strip(),
                    metadata={
                        "source": filename,
                        "section": current_heading or "Intro",
                        "file_type": "markdown",
                    },
                )
            )
            buf = []

    for line in lines:
        stripped = line.lstrip()
        if stripped.startswith("#"):
            flush()
            current_heading = stripped.lstrip("#").strip()
        elif line.strip():
            buf.append(line.rstrip())
    flush()
    return docs


def parse_file(path: str, filename: str) -> list[Document]:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    if ext == ".pdf":
        docs = parse_pdf(path, filename)
    elif ext == ".docx":
        docs = parse_docx(path, filename)
    else:  # .md / .markdown / .txt
        docs = parse_markdown(path, filename)

    if not docs:
        raise EmptyDocument(f"No extractable text found in '{filename}'.")
    return docs
